#!/usr/bin/env python3
"""place-scan 리포트(Markdown)를 MS Word(.docx)로 변환한다.

사용법:
    python tools/md_to_docx.py input.md output.docx

지원: 제목(#/##/###), 표(| a | b |), 불릿(- ), **굵게**, `코드/배지`,
[텍스트](url)/<url> 링크, 수평선(---). python-docx만 사용한다.
클라우드 루틴에서 python-docx가 없으면 먼저 설치한다:
    python -m pip install --quiet python-docx
"""
import re
import sys


def _add_runs(paragraph, text):
    """인라인 마크업(**굵게**, `코드`, [t](u), <u>)을 런으로 추가."""
    # 링크를 표시 텍스트로 치환하며 url을 괄호로 남긴다
    text = re.sub(r"\[([^\]]+)\]\((https?://[^)]+)\)", r"\1 (\2)", text)
    text = re.sub(r"<(https?://[^>]+)>", r"\1", text)
    # **굵게** 와 `코드` 를 토큰으로 분리
    for token in re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        else:
            paragraph.add_run(token)


def _is_table_row(line):
    return line.strip().startswith("|") and line.strip().endswith("|")


def _split_row(line):
    cells = line.strip().strip("|").split("|")
    return [c.strip() for c in cells]


def convert(md_path, docx_path):
    from docx import Document
    from docx.shared import Pt

    with open(md_path, encoding="utf-8") as f:
        lines = f.read().splitlines()

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Malgun Gothic"
    style.font.size = Pt(11)

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 표: 헤더 줄 + 구분줄(---) + 본문 줄
        if _is_table_row(line) and i + 1 < len(lines) and set(lines[i + 1].strip()) <= set("|-: "):
            header = _split_row(line)
            rows = []
            i += 2  # 헤더줄 + 구분줄 건너뜀
            while i < len(lines) and _is_table_row(lines[i]):
                rows.append(_split_row(lines[i]))
                i += 1
            table = doc.add_table(rows=1, cols=len(header))
            table.style = "Light Grid Accent 1"
            for c, htext in enumerate(header):
                cell = table.rows[0].cells[c]
                cell.paragraphs[0].clear()
                _add_runs(cell.paragraphs[0], htext)
                for run in cell.paragraphs[0].runs:
                    run.bold = True
            for row in rows:
                cells = table.add_row().cells
                for c, val in enumerate(row[: len(header)]):
                    cells[c].paragraphs[0].clear()
                    _add_runs(cells[c].paragraphs[0], val)
            doc.add_paragraph("")
            continue

        if not stripped:
            i += 1
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif set(stripped) <= set("-") and len(stripped) >= 3:
            pass  # 수평선은 생략
        elif stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            _add_runs(p, stripped[2:])
        else:
            p = doc.add_paragraph()
            _add_runs(p, stripped)
        i += 1

    doc.save(docx_path)


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("usage: python tools/md_to_docx.py input.md output.docx", file=sys.stderr)
        sys.exit(2)
    convert(sys.argv[1], sys.argv[2])
    print(f"wrote {sys.argv[2]}")
